from docx import Document
from docx.text.run import Font
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm


doc = Document()

doc.add_heading("Judul Laporan", 0)

table = doc.add_table(rows=1, cols=5)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = "Nama Kayu"
hdr_cells[1].text = "Jalur"
hdr_cells[2].text = "Diameter Setinggi Dada(CM)"
hdr_cells[3].text = "Tinggi Bebas Cabang(M)"
hdr_cells[4].text = "Volume Pohon(M^3)"


def func(nm_kayu, a, b, c, d):
    row_cells = table.add_row().cells
    row_cells[0].text = nm_kayu
    row_cells[1].text = str(a)
    row_cells[2].text = str(b)
    row_cells[3].text = str(c)
    row_cells[4].text = str(d)

    doc.save(f"laporan/laporan_kayu_{a}.docx")
